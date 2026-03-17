from pathlib import Path

from docx import Document

base = Path(__file__).resolve().parents[1]
md_path = base / "README_COMPARTILHAR.md"
out_path = base / "README_COMPARTILHAR.docx"

text = md_path.read_text(encoding="utf-8").splitlines()
doc = Document()

for raw in text:
    line = raw.rstrip()
    if not line:
        doc.add_paragraph("")
        continue

    if line.startswith("### "):
        doc.add_heading(line[4:].strip(), level=3)
    elif line.startswith("## "):
        doc.add_heading(line[3:].strip(), level=2)
    elif line.startswith("# "):
        doc.add_heading(line[2:].strip(), level=1)
    elif line.startswith("- "):
        doc.add_paragraph(line[2:].strip(), style="List Bullet")
    elif line[:3].isdigit() and line[1:3] == ". ":
        doc.add_paragraph(line[3:].strip(), style="List Number")
    elif line.startswith("```"):
        # Ignora marcadores de bloco para manter o texto limpo no Word.
        continue
    else:
        doc.add_paragraph(line)

doc.save(out_path)
print(f"Arquivo gerado: {out_path}")
